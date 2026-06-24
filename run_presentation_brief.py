"""
run_presentation_brief.py — PT Research Pipeline v2 — Presentation Brief

Generates a structured Word (.docx) clinical evidence brief after each pipeline run.
The document uses proper Word heading hierarchy (H1/H2/H3) so Copilot, Word's
Document Map, and human readers can navigate it cleanly.

Document structure:
  1. Cover — topic, date, pipeline version, N summary
  2. Part I: Evidence Landscape — ingested N, screening funnel, Oxford distribution
     (L0 grades), GRADE downgrade note, post-appraisal Oxford distribution (L2 grades)
  3. Part II: Cluster Summaries — one H2 section per cluster with evidence spectrum
     indicator, key findings, clinician recommendation, governance gap
  4. Part III: Clinical Impact (if Stage 5b ran) — leverage rubric, then one H2
     per subclassification with top intervention, MCID status, protocol
  5. Part IV: Spin & Governance Overview
  6. References — AMA style, numbered, matching in-text superscript citations

Usage:
    python3 run_presentation_brief.py
    python3 run_presentation_brief.py --output my_brief.docx
    python3 run_presentation_brief.py --no-clinical-impact
"""

import argparse
import json
import logging
import os
import re
import sys
import yaml
import pandas as pd
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

PIPELINE_ROOT = Path(__file__).parent.resolve()
sys.path.insert(0, str(PIPELINE_ROOT))
os.chdir(PIPELINE_ROOT)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] — %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
log = logging.getLogger(__name__)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x19, 0x23)
TEAL   = RGBColor(0x2A, 0x9D, 0x8F)
TEAL2  = RGBColor(0x1A, 0x7A, 0x6E)
GOLD   = RGBColor(0xE9, 0xC4, 0x6A)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)

# ── Oxford conversion ─────────────────────────────────────────────────────────
OXFORD_TO_ROMAN = {
    "1a": "I", "1b": "I", "1c": "I",
    "2a": "II", "2b": "II", "2c": "II",
    "3a": "III", "3b": "III",
    "4": "IV", "5": "V",
    "I": "I", "II": "II", "III": "III", "IV": "IV", "V": "V",
}
ROMAN_ORDER = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "?": 6}
OXFORD_DESCRIPTIONS = {
    "I":   "Systematic reviews, meta-analyses, and high-quality RCTs",
    "II":  "Prospective cohort studies and lower-quality RCTs",
    "III": "Retrospective cohort or case-control studies",
    "IV":  "Case series and poor-quality cohort/case-control studies",
    "V":   "Expert opinion, narrative reviews, and consensus statements",
}
GRADE_ICON = {
    "High": "●●●●", "Moderate": "●●●○", "Low": "●●○○", "Very Low": "●○○○"
}

# Words that should stay uppercase regardless of position
ALWAYS_UPPER = {"pt", "ai", "rct", "bppv", "pfmt", "uti", "lbp", "copd",
                "acl", "mcl", "pcl", "mri", "us", "emg", "eswt", "tens",
                "nmes", "tms", "vr", "app"}
# Words that should stay lowercase (articles/prepositions) unless first word
ALWAYS_LOWER = {"and", "or", "of", "the", "in", "for", "with", "a", "an",
                "to", "at", "by", "vs", "per"}


def clean_label(raw: str) -> str:
    """Convert any taxonomy ID or raw string to a clean human-readable label.

    Handles patterns like:
      PF_stress_urinary_incontinence  → Stress Urinary Incontinence
      LBP_radiculopathy               → Radiculopathy
      CANINE_ankle_fracture_rehab     → Ankle Fracture Rehab
      fecal_incontinence_disorders    → Fecal Incontinence Disorders
      Pf Athletic Special             → Athletic Special  (strips rendered prefix)
      Female Sexual Dysfunction       → Female Sexual Dysfunction (passthrough)
    """
    if not raw:
        return ""
    import re as _re
    s = str(raw)
    # Strip uppercase namespace prefix: e.g. PF_, LBP_, CANINE_
    s = _re.sub(r'^[A-Z]{1,8}_', '', s)
    # Strip already-rendered taxonomy prefix words (e.g. "Pf ", "Lbp ", "Bppv ")
    # These appear when taxonomy IDs were .title()-d before storage.
    # Only strip if the first word is a known short abbreviation-style prefix.
    RENDERED_PREFIXES = {"pf", "lbp", "bppv", "dn", "vr", "sui", "oab",
                         "ui", "uti", "msk", "tbi", "cvd", "gi", "oa"}
    words_preview = s.split()
    if (len(words_preview) > 1 and
            words_preview[0].lower() in RENDERED_PREFIXES and
            words_preview[1][0].isupper()):
        s = " ".join(words_preview[1:])
    # Replace underscores/hyphens with spaces
    s = s.replace("_", " ").replace("-", " ")
    # Title-case word by word, respecting abbreviations
    words = s.split()
    result = []
    for i, word in enumerate(words):
        lower = word.lower()
        if lower in ALWAYS_UPPER:
            result.append(word.upper())
        elif lower in ALWAYS_LOWER and i > 0:
            result.append(lower)
        else:
            result.append(word.capitalize())
    return " ".join(result)
GRADE_COLOR = {
    "High": TEAL, "Moderate": TEAL2, "Low": GOLD, "Very Low": RED
}
GRADE_ORDER = {"High": 0, "Moderate": 1, "Low": 2, "Very Low": 3}

EVIDENCE_SPECTRUM = [
    ("Overwhelming",  9, 10, TEAL),
    ("Strong",        7,  8, TEAL2),
    ("Moderate",      5,  6, RGBColor(0x2E, 0x86, 0xAB)),
    ("Limited",       3,  4, NAVY),
    ("Insufficient",  0,  2, GREY),
]

def score_to_label(score):
    try:
        s = int(score)
        for label, lo, hi, _ in EVIDENCE_SPECTRUM:
            if lo <= s <= hi:
                return label, s
    except (TypeError, ValueError):
        pass
    return "Unscored", None


def to_roman(raw: str) -> str:
    if not raw:
        return "?"
    return OXFORD_TO_ROMAN.get(str(raw).strip(), "?")


def oxford_distribution(articles: list, field="oxford_roman") -> dict:
    counts = Counter()
    for a in articles:
        raw = a.get(field, "") or a.get("oxford_level", "")
        counts[to_roman(str(raw))] += 1
    return counts


# ── Citation manager ──────────────────────────────────────────────────────────
class CitationManager:
    """Assigns sequential numbers to PMIDs, builds AMA reference list."""
    def __init__(self):
        self._map: dict[str, int] = {}
        self._articles: dict[str, dict] = {}

    def register(self, article: dict) -> int:
        pmid = str(article.get("pmid", ""))
        if not pmid:
            return 0
        if pmid not in self._map:
            self._map[pmid] = len(self._map) + 1
            self._articles[pmid] = article
        return self._map[pmid]

    def number(self, pmid: str) -> int:
        return self._map.get(str(pmid), 0)

    def references(self) -> list[tuple[int, str]]:
        """Return sorted list of (number, AMA_formatted_citation)."""
        refs = []
        for pmid, num in sorted(self._map.items(), key=lambda x: x[1]):
            a = self._articles[pmid]
            authors_raw = a.get("authors", "") or ""
            if isinstance(authors_raw, list):
                authors_raw = "; ".join(authors_raw)
            # Format authors: Last FM style, max 3 then et al
            parts = [p.strip() for p in re.split(r";|,\s+(?=[A-Z])", authors_raw) if p.strip()]
            if len(parts) > 3:
                author_str = f"{parts[0]}, et al"
            elif parts:
                author_str = ", ".join(parts)
            else:
                author_str = "Unknown"
            title   = (a.get("title", "") or "").strip().rstrip(".")
            journal = (a.get("journal", "") or "").strip()
            year    = str(a.get("year", "") or "").strip()
            doi     = (a.get("doi", "") or "").strip()
            pmid_str = pmid
            # AMA format: Author(s). Title. Journal. Year. doi:xxx. PMID:xxx
            ref = f"{author_str}. {title}. {journal}. {year}."
            if doi:
                ref += f" doi:{doi}."
            if pmid_str:
                ref += f" PMID:{pmid_str}."
            refs.append((num, ref))
        return refs


# ── python-docx helpers ────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_superscript(run):
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement("w:vertAlign")
    vertAlign.set(qn("w:val"), "superscript")
    rPr.append(vertAlign)


def add_run(para, text: str, bold=False, italic=False, size=None,
            color=None, superscript=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if superscript:
        add_superscript(run)
    return run


def para_after(doc, text="", bold=False, italic=False, size=None,
               color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
               space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.color.rgb = TEAL
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.color.rgb = NAVY
    return p


def bullet(doc, text: str, level=0, bold_prefix: str = None, citations: list = None,
           cite_mgr: CitationManager = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=11)
        add_run(p, text, size=11)
    else:
        add_run(p, text, size=11)
    if citations and cite_mgr:
        for pmid in citations:
            num = cite_mgr.number(str(pmid))
            if num:
                add_run(p, str(num), size=8, superscript=True)
    return p


def add_divider(doc, color_hex="2A9D8F"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    # pBdr must appear before spacing/ind/jc in pPr per OOXML schema
    # Insert at position 0 (before any spacing elements)
    pPr.insert(0, pBdr)
    return p


def evidence_spectrum_table(doc, clusters_with_scores: list):
    """Vertical spectrum table — full cluster names in left column, band in right.
    Sorted strongest evidence first. No truncation."""
    if not clusters_with_scores:
        return

    BAND_COLORS = {
        "Overwhelming":  "1A7A6E",
        "Strong":        "2A9D8F",
        "Moderate":      "2E86AB",
        "Limited":       "0F1923",
        "Insufficient":  "888888",
    }
    BAND_ORDER = {"Overwhelming": 0, "Strong": 1, "Moderate": 2,
                  "Limited": 3, "Insufficient": 4, "Unscored": 5}

    # Sort by evidence strength
    rows = []
    for cluster_label, score in clusters_with_scores:
        label, _ = score_to_label(score)
        score_val = score if isinstance(score, (int, float)) else 0
        rows.append((cluster_label, score_val, label))
    rows.sort(key=lambda x: BAND_ORDER.get(x[2], 5))

    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for ci, (hdr, w) in enumerate(
            [("Clinical Cluster", 4.5), ("Score", 0.8), ("Evidence Band", 1.7)]):
        cell = tbl.cell(0, ci)
        set_cell_bg(cell, "0F1923")
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for ri, (cluster_label, score_val, band_label) in enumerate(rows):
        color_hex = BAND_COLORS.get(band_label, "888888")
        bg = "F0F6F8" if ri % 2 == 0 else "FFFFFF"

        # Cluster name
        cell = tbl.cell(ri + 1, 0)
        set_cell_bg(cell, bg)
        run = cell.paragraphs[0].add_run(cluster_label)
        run.font.size = Pt(10)

        # Score
        cell = tbl.cell(ri + 1, 1)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{score_val}/10" if score_val else "—")
        run.font.size = Pt(10)
        run.font.bold = True

        # Band — coloured background
        cell = tbl.cell(ri + 1, 2)
        set_cell_bg(cell, color_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(band_label)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE

    doc.add_paragraph()


def oxford_dist_table(doc, counts: Counter, total: int, label: str):
    """Simple 5-row Oxford distribution summary table."""
    h3(doc, label)
    tbl = doc.add_table(rows=6, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Oxford Level", "Description", "n (%)"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_bg(cell, "0F1923")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    roman_levels = ["I", "II", "III", "IV", "V"]
    for ri, level in enumerate(roman_levels):
        n = counts.get(level, 0)
        pct = f"{n/total*100:.0f}%" if total else "0%"
        bg = "F0F6F8" if ri % 2 == 0 else "FFFFFF"
        row_data = [level, OXFORD_DESCRIPTIONS.get(level, ""), f"{n} ({pct})"]
        for ci, text in enumerate(row_data):
            cell = tbl.cell(ri + 1, ci)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            if ci == 0:
                run.font.bold = True
    doc.add_paragraph()


def grade_spectrum_bar(doc, cert: str):
    """Visual GRADE certainty indicator as a simple text bar."""
    icon = GRADE_ICON.get(cert, "○○○○")
    color = GRADE_COLOR.get(cert, GREY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "GRADE Certainty: ", bold=True, size=11)
    add_run(p, f"{icon} {cert}", bold=True, size=11, color=color)
    return p


# ── Main builder ──────────────────────────────────────────────────────────────
def build_brief(cfg: dict, articles: list, syntheses: list,
                cluster_defs: dict, ci_df, ci_cfg: dict,
                cite_mgr: CitationManager) -> Document:

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    topic     = cfg.get("topic", {}).get("short_name", "Evidence Synthesis")
    # Preserve original casing (e.g. "pelvic floor PT" not "Pelvic Floor Pt")
    topic_display = " ".join(
        w.upper() if w.lower() in ("pt", "ai", "rct", "bppv") else w.capitalize()
        for w in topic.split()
    )
    rq        = cfg.get("research_question", "").strip()
    today     = datetime.now().strftime("%B %d, %Y")
    n_total   = len(articles)
    threshold = ci_cfg.get("leverage_threshold", 7) if ci_cfg else 7

    # ── Layer 0 data ──────────────────────────────────────────────────────────
    l0_path = Path("layer0_ledger.csv")
    l0_df   = pd.read_csv(l0_path) if l0_path.exists() else pd.DataFrame()
    n_screened = len(l0_df) if not l0_df.empty else "N/A"
    n_relevant = int(l0_df["relevant_to_primary_question"].str.lower().eq("yes").sum()) \
                 if not l0_df.empty and "relevant_to_primary_question" in l0_df.columns else n_total
    n_excluded = int(n_screened) - n_relevant if isinstance(n_screened, int) else "N/A"

    # Register all articles for citation
    for a in articles:
        cite_mgr.register(a)

    # ── COVER ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    add_run(p, f"Evidence Brief: {topic_display}", bold=True, size=22, color=NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "PT Research Pipeline v2 — Automated Evidence Synthesis", size=12, color=GREY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p3, f"Generated: {today}  |  N = {n_total} articles appraised  |  "
                f"{len(syntheses)} clinical clusters identified",
            size=11, color=GREY)

    add_divider(doc)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_after = Pt(6)
    add_run(p4, "Research Question: ", bold=True, size=11)
    add_run(p4, rq, italic=True, size=11, color=GREY)

    doc.add_page_break()

    # ── PART I: EVIDENCE LANDSCAPE ────────────────────────────────────────────
    h1(doc, "Part I: Evidence Landscape")
    add_divider(doc)

    h2(doc, "1.1 Screening Funnel")
    bullet(doc, f"Total records identified via PubMed RSS / E-utilities: {n_screened}")
    bullet(doc, f"Excluded as irrelevant to research question: {n_excluded}")
    bullet(doc, f"Relevant articles forwarded to deep appraisal: {n_relevant}")
    bullet(doc, f"Articles completing full 7-stage appraisal (Layer 2): {n_total}")
    doc.add_paragraph()

    # Layer 0 Oxford distribution
    if not l0_df.empty and "oxford_level" in l0_df.columns:
        l0_counts = Counter()
        for val in l0_df["oxford_level"].dropna():
            l0_counts[to_roman(str(val))] += 1
        oxford_dist_table(doc, l0_counts, int(n_screened),
                          "1.2 Oxford Evidence Level Distribution — Screened Articles (Layer 0)")

    # Layer 2 Oxford distribution (post-appraisal, potentially downgraded)
    l2_counts = oxford_distribution(articles, "oxford_roman")
    oxford_dist_table(doc, l2_counts, n_total,
                      "1.3 Oxford Evidence Level Distribution — Appraised Articles (Layer 2)")

    p = doc.add_paragraph()
    add_run(p, "Note on evidence level downgrading: ", bold=True, size=11)
    add_run(p, "Oxford levels assigned during Layer 0 screening are based on study design alone. "
               "During Layer 2 deep appraisal, the model applies quality adjustment — "
               "accounting for sample size, blinding, allocation concealment, dropout rate, "
               "and other methodological factors. Studies are frequently downgraded from their "
               "initial design-based level. The Layer 2 distribution above reflects "
               "quality-adjusted evidence grades.", size=11)

    # GRADE distribution
    h2(doc, "1.4 GRADE Certainty Distribution Across Clusters")
    if syntheses:
        grade_counts = Counter(s.get("grade_certainty", "?") for s in syntheses)
        for grade in ["High", "Moderate", "Low", "Very Low"]:
            n = grade_counts.get(grade, 0)
            icon = GRADE_ICON.get(grade, "")
            color = GRADE_COLOR.get(grade, GREY)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_run(p, f"{icon} {grade}: ", bold=True, size=11, color=color)
            add_run(p, f"{n} cluster(s) ({n/len(syntheses)*100:.0f}%)", size=11)
    doc.add_paragraph()

    spin_n   = sum(1 for a in articles if (a.get("spin_detected","") or "").lower() == "yes")
    spin_pct = f"{spin_n/n_total*100:.0f}%" if n_total else "0%"
    h2(doc, "1.5 Spin Detection Summary")
    bullet(doc, f"{spin_n} of {n_total} articles ({spin_pct}) exhibited detectable spin — "
                "conclusion language that overstated what the reported numbers support.")
    bullet(doc, "Spin types detected: Causal Leap, Abstract Disconnect, Suppression, "
                "False Precision, Protocol Claiming Results.")
    bullet(doc, "Spin detection is automated; all flagged articles should be verified by "
                "a qualified clinician before clinical policy decisions.")

    doc.add_page_break()

    # ── PART II: CLUSTER SUMMARIES ────────────────────────────────────────────
    h1(doc, "Part II: Clinical Cluster Summaries")
    add_divider(doc)

    # Evidence spectrum overview table
    h2(doc, "2.1 Evidence Strength Spectrum — All Clusters")
    para_after(doc, "The table below maps each cluster to its GRADE certainty band. "
                    "Clusters to the left have stronger, more consistent evidence. "
                    "Clusters to the right have limited, heterogeneous, or insufficient evidence.",
               size=11)

    # Build cluster score list for spectrum (use GRADE order as proxy score)
    grade_to_score = {"High": 9, "Moderate": 7, "Low": 5, "Very Low": 3}
    cluster_scores = []
    for s in syntheses:
        label = clean_label(s.get("cluster", ""))
        cert  = s.get("grade_certainty", "Very Low")
        score = grade_to_score.get(cert, 3)
        cluster_scores.append((label, score))
    evidence_spectrum_table(doc, cluster_scores)

    # Per-cluster sections
    sorted_synths = sorted(syntheses,
                           key=lambda s: GRADE_ORDER.get(s.get("grade_certainty",""), 9))

    for idx, s in enumerate(sorted_synths, 1):
        cluster_name  = s.get("cluster", "")
        cluster_label = clean_label(cluster_name)
        cert          = s.get("grade_certainty", "Very Low")
        rec_dir       = s.get("recommendation_direction", "").upper()
        rec_str       = s.get("recommendation_strength", "").title()
        key_findings  = s.get("key_findings", "")
        clin_rec      = s.get("clinician_recommendation", "")
        caveats       = s.get("key_caveat", "")
        gov_rec       = s.get("governance_recommendation", "")
        future_res    = s.get("future_research_priority", "")

        # Cluster articles — match by cluster field, fall back to synthesis n_articles
        cluster_articles = [a for a in articles
                            if a.get("cluster", "") == cluster_name]
        n_cluster = len(cluster_articles)
        # Fallback: use n_articles from the synthesis JSON if available and matching gives 0
        if n_cluster == 0:
            n_cluster = s.get("n_articles", s.get("article_count", 0)) or 0

        h2(doc, f"2.{idx}  {cluster_label}")
        grade_spectrum_bar(doc, cert)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run(p, "Recommendation: ", bold=True, size=11)
        color = TEAL if rec_dir == "FOR" else RED if rec_dir == "AGAINST" else GREY
        add_run(p, f"{rec_dir} ({rec_str})", bold=True, size=11, color=color)
        add_run(p, f"   |   Articles: {n_cluster}", size=11, color=GREY)

        h3(doc, "Key Findings")
        if key_findings:
            # Split compound findings into bullets and add citations
            sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", key_findings) if s.strip()]
            for sent in sentences:
                # Find PMIDs mentioned via citation manager
                cites = [a.get("pmid") for a in cluster_articles[:3]]
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                add_run(p, sent, size=11)
                for pmid in cites:
                    if pmid:
                        num = cite_mgr.number(str(pmid))
                        if num:
                            add_run(p, str(num), size=8, superscript=True)

        h3(doc, "Clinician Recommendation")
        if clin_rec:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            add_run(p, clin_rec, italic=True, size=11)

        if caveats:
            h3(doc, "Key Caveat")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_run(p, caveats, size=11, color=GREY)

        if gov_rec:
            h3(doc, "Governance Gap")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_run(p, gov_rec, size=11)

        if future_res:
            h3(doc, "Future Research Priority")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_run(p, future_res, size=11, color=GREY)

        add_divider(doc, "CCCCCC")

    doc.add_page_break()

    # ── PART III: CLINICAL IMPACT ─────────────────────────────────────────────
    ci_enabled = (ci_df is not None and not ci_df.empty)

    h1(doc, "Part III: Clinical Impact Analysis")
    add_divider(doc)

    if not ci_enabled:
        para_after(doc, "Clinical Impact (Stage 5b) was not run for this topic. "
                        "To enable, set clinical_impact.enabled: true in config.yaml "
                        "and re-run the pipeline.", italic=True, size=11, color=GREY)
    else:
        # Leverage rubric
        h2(doc, "3.1 Clinical Leverage Score — Rubric")
        para_after(doc, "Each article is scored from 0–10 against the following criteria. "
                        f"A score of {threshold} or above triggers a Maximal Impact Flag.",
                   size=11)
        rubric = [
            ("+3", "Between-group difference meets or exceeds MCID",
             "Did the treatment move the needle enough to matter to the patient?"),
            ("+2", "Effect size ≥ 0.8 (large)",
             "Cohen's d or Hedges' g indicates a clinically large effect"),
            ("+1", "Primary outcome p < 0.05", "Statistically significant result"),
            ("+1", "Protocol fully specified",
             "Dose, frequency, and duration all stated and replicable"),
            ("+1", "Low or moderate bias risk",
             "Study design and execution meet quality thresholds"),
            ("+1", "RCT or higher study design", "Oxford level Ia, Ib, or IIa"),
            ("+1", "No spin detected",
             "Conclusions accurately reflect what the numbers show"),
        ]
        tbl = doc.add_table(rows=len(rubric)+1, cols=3)
        tbl.style = "Table Grid"
        headers = ["Points", "Criterion", "Meaning"]
        for ci, hdr in enumerate(headers):
            cell = tbl.cell(0, ci)
            set_cell_bg(cell, "0F1923")
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
        for ri, (pts, crit, meaning) in enumerate(rubric):
            bg = "E8F5F3" if ri % 2 == 0 else "FFFFFF"
            for ci, text in enumerate([pts, crit, meaning]):
                cell = tbl.cell(ri+1, ci)
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(10)
                if ci == 0:
                    run.font.bold = True
                    run.font.color.rgb = TEAL if pts in ("+3", "+2") else NAVY
        doc.add_paragraph()

        # Star scale
        p = doc.add_paragraph()
        add_run(p, "Score interpretation: ", bold=True, size=11)
        scale = [("9–10 ★★★★★ Overwhelming", TEAL),
                 ("  7–8 ★★★★ Strong", TEAL2),
                 ("  5–6 ★★★ Moderate", RGBColor(0x2E, 0x86, 0xAB)),
                 ("  3–4 ★★ Limited", NAVY),
                 ("  0–2 ★ Insufficient", GREY)]
        for text, color in scale:
            add_run(p, text, size=10, color=color, bold=True)
        doc.add_paragraph()

        # Per-subclassification sections
        h2(doc, "3.2 Intervention Hierarchy by Clinical Subclassification")

        conditions = ci_df["condition_classification"].dropna().unique() \
                     if "condition_classification" in ci_df.columns else []

        CONTROL_PATS = ["control", "sham", "placebo", "standard care",
                        "usual care", "waitlist", "no treatment", "not_applicable"]
        def is_ctrl(arm):
            if not arm or str(arm) in ("not_reported", "nan"):
                return True
            return any(p in str(arm).lower() for p in CONTROL_PATS)

        for condition in sorted(conditions):
            cdf = ci_df[ci_df["condition_classification"] == condition]
            n_articles = cdf["pmid"].nunique() if "pmid" in cdf.columns else len(cdf)
            med_score = pd.to_numeric(
                cdf.drop_duplicates("pmid")["clinical_leverage_score"]
                if "pmid" in cdf.columns else cdf["clinical_leverage_score"],
                errors="coerce"
            ).median()
            label, _ = score_to_label(med_score)

            h3(doc, clean_label(condition))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            add_run(p, f"Articles: {n_articles}   |   Median Leverage: ", size=11)
            score_str = f"{med_score:.1f}/10" if pd.notna(med_score) else "—"
            add_run(p, score_str, bold=True, size=11, color=TEAL if pd.notna(med_score) and med_score >= 7 else GREY)
            add_run(p, f"   |   {label}", size=11, color=GREY)

            # Max impact alerts
            if "clinical_leverage_score" in cdf.columns:
                hi = cdf[pd.to_numeric(cdf["clinical_leverage_score"], errors="coerce").fillna(0) >= threshold]
                hi_articles = hi.drop_duplicates("pmid") if "pmid" in hi.columns else hi
                for _, row in hi_articles.iterrows():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_after = Pt(3)
                    add_run(p, f"🔴 MAXIMAL IMPACT  Score {row.get('clinical_leverage_score','?')}/10  — ",
                            bold=True, size=11, color=RED)
                    headline = row.get("alert_headline", "")
                    if headline:
                        add_run(p, str(headline), size=11)
                    # Citation
                    pmid = str(row.get("pmid",""))
                    if pmid and cite_mgr.number(pmid):
                        add_run(p, str(cite_mgr.number(pmid)), size=8, superscript=True)

            # Intervention hierarchy
            if "winning_arm" in cdf.columns:
                hierarchy = []
                for arm, grp in cdf.groupby("winning_arm", dropna=True):
                    if is_ctrl(arm):
                        continue
                    scores = pd.to_numeric(grp["clinical_leverage_score"], errors="coerce").dropna()
                    med = round(scores.median(), 1) if len(scores) else None
                    mcid_n = grp["mcid_met"].str.lower().isin(["yes","borderline"]).sum() \
                             if "mcid_met" in grp.columns else 0
                    proto = ""
                    if "winning_arm_protocol" in grp.columns:
                        modes = grp["winning_arm_protocol"].dropna()
                        proto = str(modes.mode()[0])[:120] if len(modes) else ""
                    hierarchy.append((arm, med, len(grp), mcid_n, proto))
                hierarchy.sort(key=lambda x: x[1] or 0, reverse=True)

                if hierarchy:
                    tbl = doc.add_table(rows=len(hierarchy)+1, cols=4)
                    tbl.style = "Table Grid"
                    for ci2, hdr in enumerate(["Intervention","Leverage","MCID Met","Protocol (summary)"]):
                        cell = tbl.cell(0, ci2)
                        set_cell_bg(cell, "162230")
                        p = cell.paragraphs[0]
                        run = p.add_run(hdr)
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = WHITE
                    for ri2, (arm, med, n_comp, mcid_n, proto) in enumerate(hierarchy):
                        bg = "E8F5F3" if ri2 == 0 else ("F5F5F5" if ri2 % 2 == 0 else "FFFFFF")
                        score_txt = f"{med}/10" if med is not None else "—"
                        star_label, _ = score_to_label(med)
                        for ci2, text in enumerate([
                            arm[:60],
                            f"{score_txt} {star_label}",
                            f"{mcid_n}/{n_comp}",
                            proto[:100] if proto else "not_reported"
                        ]):
                            cell = tbl.cell(ri2+1, ci2)
                            set_cell_bg(cell, bg)
                            p = cell.paragraphs[0]
                            run = p.add_run(text)
                            run.font.size = Pt(9)
                            if ci2 == 0 and ri2 == 0:
                                run.font.bold = True
                    doc.add_paragraph()

    doc.add_page_break()

    # ── PART IV: GOVERNANCE OVERVIEW ──────────────────────────────────────────
    h1(doc, "Part IV: Governance & Spin Overview")
    add_divider(doc)

    GOVERNANCE_DIMENSIONS = [
        ("scope_of_practice",    "Scope of Practice"),
        ("output_validation",    "Output Validation"),
        ("guardrails_safety",    "Guardrails / Safety"),
        ("accountability_liability", "Accountability / Liability"),
        ("training_competency",  "Training / Competency"),
    ]

    h2(doc, "4.1 Governance Taxonomy — Corpus-Wide")
    para_after(doc, f"Each of the {n_total} appraised articles was audited across five governance "
                    "dimensions: Implemented (authors actually did this in their study), "
                    "Aspirational (authors recommend it without implementing it), "
                    "or Not Addressed (no mention at all).", size=11)

    tbl = doc.add_table(rows=len(GOVERNANCE_DIMENSIONS)+1, cols=4)
    tbl.style = "Table Grid"
    for ci, hdr in enumerate(["Dimension","Implemented","Aspirational","Not Addressed"]):
        cell = tbl.cell(0, ci)
        set_cell_bg(cell, "0F1923")
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for ri, (dim_key, dim_label) in enumerate(GOVERNANCE_DIMENSIONS):
        impl = aspir = not_addr = 0
        for a in articles:
            dims = a.get("governance_dimensions", {}) or {}
            status = (dims.get(dim_key, {}) or {}).get("status", "not_addressed")
            if status == "implemented":   impl += 1
            elif status == "aspirational": aspir += 1
            else:                          not_addr += 1
        bg = "F0F6F8" if ri % 2 == 0 else "FFFFFF"
        is_gap = not_addr / n_total > 0.8 if n_total else False
        for ci, (text, alert) in enumerate([
            (dim_label, False),
            (f"{impl} ({impl/n_total*100:.0f}%)" if n_total else "0", False),
            (f"{aspir} ({aspir/n_total*100:.0f}%)" if n_total else "0", False),
            (f"{not_addr} ({not_addr/n_total*100:.0f}%)" if n_total else "0", is_gap),
        ]):
            cell = tbl.cell(ri+1, ci)
            set_cell_bg(cell, "FFE0DC" if alert else bg)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.bold = (ci == 0 or alert)
            if alert:
                run.font.color.rgb = RED
    doc.add_paragraph()

    doc.add_page_break()

    # ── REFERENCES ────────────────────────────────────────────────────────────
    h1(doc, "References")
    add_divider(doc)
    para_after(doc, "References are formatted in AMA style. Numbers correspond to "
                    "in-text citations throughout this document.", size=11, color=GREY)
    doc.add_paragraph()

    refs = cite_mgr.references()
    for num, ref_text in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after  = Pt(4)
        add_run(p, f"{num}.  ", bold=True, size=10, color=TEAL)
        add_run(p, ref_text, size=10)

    # Methods note
    doc.add_page_break()
    h1(doc, "Methods Note")
    add_divider(doc)
    model = cfg.get("model", {}).get("layer2", "qwen3.6:35b-a3b")
    para_after(doc, f"This evidence brief was generated by PT Research Pipeline v2 "
                    f"using {model} for all three processing layers (screening, deep appraisal, "
                    f"and cluster synthesis). Layer 0 applies Oxford OCEBM evidence levels "
                    f"based on study design. Layer 2 performs quality-adjusted regrading "
                    f"after 7-stage methodological appraisal. Layer 3 applies GRADE "
                    f"certainty of evidence ratings per dynamically discovered clinical cluster.",
               size=11)
    para_after(doc, "All AI-generated outputs should be reviewed by a qualified clinician "
                    "before informing clinical policy, practice guidelines, or patient care decisions.",
               size=11, italic=True, color=GREY)
    para_after(doc, f"github.com/crappybanjos-del/pt-research-pipeline  |  "
                    f"DOI: 10.5281/zenodo.20709834  |  Generated: {today}",
               size=10, color=GREY)

    return doc


# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="Generate structured clinical evidence brief")
    parser.add_argument("--output", default=None, help="Output .docx filename")
    parser.add_argument("--no-clinical-impact", action="store_true",
                        help="Exclude Clinical Impact section even if ledger exists")
    args = parser.parse_args()

    cfg = yaml.safe_load(Path("config.yaml").read_text())
    ci_cfg = cfg.get("clinical_impact", {})

    # Load data
    deep_dir = Path("summaries/deep")
    articles = []
    if deep_dir.exists():
        for f in sorted(deep_dir.glob("*_layer2.json")):
            try:
                articles.append(json.loads(f.read_text()))
            except Exception:
                pass

    syntheses = []
    master = Path("summaries/layer3/MASTER_GRADE_SYNTHESIS.json")
    if master.exists():
        syntheses = json.loads(master.read_text())

    cluster_defs = {}
    cdef_path = Path("layer3_cluster_definitions.json")
    if cdef_path.exists():
        cluster_defs = json.loads(cdef_path.read_text()).get("definitions", {})

    ci_df = None
    if not args.no_clinical_impact and ci_cfg.get("enabled", False):
        ledger = Path("clinical_impact_ledger.csv")
        if ledger.exists():
            ci_df = pd.read_csv(ledger)
            log.info(f"Clinical impact ledger loaded: {len(ci_df)} rows")

    if not articles:
        log.error("No Layer 2 JSON files found in summaries/deep/ — run the pipeline first.")
        sys.exit(1)

    log.info(f"Building brief: {len(articles)} articles, {len(syntheses)} clusters")

    cite_mgr = CitationManager()
    doc = build_brief(cfg, articles, syntheses, cluster_defs, ci_df, ci_cfg, cite_mgr)

    topic = cfg.get("topic", {}).get("short_name", "evidence").replace(" ", "_").lower()
    today_str = datetime.now().strftime("%Y%m%d")
    out_name = args.output or f"{topic}_presentation_brief_{today_str}.docx"
    doc.save(out_name)

    ref_count = len(cite_mgr.references())
    log.info(f"Brief saved: {out_name}")
    log.info(f"  Articles: {len(articles)}  |  Clusters: {len(syntheses)}  |  References: {ref_count}")


if __name__ == "__main__":
    main()
